from tkinter import *
from tkinter import ttk, messagebox
from tkcalendar import DateEntry
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
from controladores import obtener_historial_clinico_completo, exportar_historial_clinico_excel, eliminar_historial, agregar_historia_clinica, obtener_paciente_por_cedula, obtener_historial_clinico_por_cedula

class PantallaHistoriasClinicas:
    def __init__(self, master):
        self.master = master

    def mostrar(self):
        self.master.limpiar_contenido()

        frame_principal = Frame(self.master.frame_contenido, bg="#E0F7FA")
        frame_principal.grid(row=0, column=0, sticky="nsew")

        etiqueta_titulo = Label(frame_principal, text="Historias Clínicas", font=("Arial", 24, "bold"), bg="#00ACC1", fg="#FFFFFF")
        etiqueta_titulo.grid(row=0, column=0, pady=20, padx=20, columnspan=2, sticky="ew")

        frame_busqueda = Frame(frame_principal, bg="#E0F7FA")
        frame_busqueda.grid(row=1, column=0, sticky="ew", padx=20, pady=10)
        
        Label(frame_busqueda, text="Buscar por Cédula:", bg="#E0F7FA", fg="#333333", font=("Arial", 14)).pack(side=LEFT, padx=5)
        self.cedula_buscar_entry = Entry(frame_busqueda, font=("Arial", 14), bg="#FFFFFF", fg="#333333")
        self.cedula_buscar_entry.pack(side=LEFT, padx=5)
        Button(frame_busqueda, text="Buscar", command=self.buscar_historial, bg="#007BFF", fg="white", font=("Arial", 14, "bold"), relief=RAISED, bd=5).pack(side=LEFT, padx=5)
        Button(frame_busqueda, text="Mostrar Todo", command=self.mostrar_todo, bg="#007BFF", fg="white", font=("Arial", 14, "bold"), relief=RAISED, bd=5).pack(side=LEFT, padx=5)

        frame_tree = Frame(frame_principal, bg="#E0F7FA")
        frame_tree.grid(row=2, column=0, sticky="nsew", padx=20, pady=10)
        frame_tree.grid_rowconfigure(0, weight=1)
        frame_tree.grid_columnconfigure(0, weight=1)

        tree_scrollbar = Scrollbar(frame_tree, orient=VERTICAL)
        tree_scrollbar.grid(row=0, column=1, sticky='ns')

        self.tree_historial_clinico = ttk.Treeview(frame_tree, columns=("Cedula", "Paciente", "Fecha", "Motivo Consulta"), show='headings', yscrollcommand=tree_scrollbar.set)
        self.tree_historial_clinico.heading("Cedula", text="Cédula")
        self.tree_historial_clinico.heading("Paciente", text="Paciente")
        self.tree_historial_clinico.heading("Fecha", text="Fecha")
        self.tree_historial_clinico.heading("Motivo Consulta", text="Motivo Consulta")
        self.tree_historial_clinico.tag_configure('oddrow', background='lightgrey')
        self.tree_historial_clinico.tag_configure('evenrow', background='white')
        self.tree_historial_clinico.grid(row=0, column=0, sticky='nsew')

        tree_scrollbar.config(command=self.tree_historial_clinico.yview)

        self.cargar_datos()

        frame_botones = Frame(frame_principal, bg="#E0F7FA")
        frame_botones.grid(row=3, column=0, pady=10, columnspan=2, sticky="ew")

        boton_crear = Button(frame_botones, text="Crear Historial", command=self.crear_historial, bg="#28A745", fg="white", font=("Arial", 16, "bold"), relief=RAISED, bd=5)
        boton_crear.grid(row=0, column=0, padx=10)

        boton_eliminar = Button(frame_botones, text="Eliminar Historial", command=self.eliminar_historial, bg="#DC3545", fg="white", font=("Arial", 16, "bold"), relief=RAISED, bd=5)
        boton_eliminar.grid(row=0, column=1, padx=10)

        boton_exportar_todo = Button(frame_botones, text="Exportar Historias", command=self.exportar_historial, bg="#007BFF", fg="white", font=("Arial", 16, "bold"), relief=RAISED, bd=5)
        boton_exportar_todo.grid(row=0, column=2, padx=10)

        boton_exportar_seleccionado = Button(frame_botones, text="Exportar Seleccionado", command=self.exportar_historial_seleccionado, bg="#007BFF", fg="white", font=("Arial", 16, "bold"), relief=RAISED, bd=5)
        boton_exportar_seleccionado.grid(row=0, column=3, padx=10)

        boton_exportar_word = Button(frame_botones, text="Exportar a Word", command=self.exportar_historial_seleccionado_word, bg="#007BFF", fg="white", font=("Arial", 16, "bold"), relief=RAISED, bd=5)
        boton_exportar_word.grid(row=0, column=4, padx=10)

        boton_volver = Button(frame_botones, text="Volver", command=self.master.crear_bienvenida, bg="#FFC107", fg="white", font=("Arial", 16, "bold"), relief=RAISED, bd=5)
        boton_volver.grid(row=0, column=5, padx=10)

        # Configurar peso para redimensionamiento
        self.master.frame_contenido.grid_rowconfigure(0, weight=1)
        self.master.frame_contenido.grid_columnconfigure(0, weight=1)
        frame_principal.grid_rowconfigure(2, weight=1)
        frame_principal.grid_columnconfigure(0, weight=1)

    def cargar_datos(self):
        for item in self.tree_historial_clinico.get_children():
            self.tree_historial_clinico.delete(item)
        historial = obtener_historial_clinico_completo()
        for i, h in enumerate(historial):
            tag = 'evenrow' if i % 2 == 0 else 'oddrow'
            self.tree_historial_clinico.insert("", "end", values=(h[1], f"{h[2]} {h[3]}", h[4], h[5]), tags=(tag,))

    def buscar_historial(self):
        cedula = self.cedula_buscar_entry.get()
        if not cedula:
            messagebox.showwarning("Advertencia", "Ingrese una cédula para buscar")
            return
        historial = obtener_historial_clinico_por_cedula(cedula)
        if historial:
            for item in self.tree_historial_clinico.get_children():
                self.tree_historial_clinico.delete(item)
            for i, h in enumerate(historial):
                tag = 'evenrow' if i % 2 == 0 else 'oddrow'
                self.tree_historial_clinico.insert("", "end", values=(h[1], f"{h[2]} {h[3]}", h[4], h[5]), tags=(tag,))
        else:
            messagebox.showinfo("Información", "No se encontraron historiales para la cédula ingresada")

    def mostrar_todo(self):
        self.cargar_datos()

    def crear_historial(self):
        self.master.limpiar_contenido()

        canvas = Canvas(self.master.frame_contenido, bg="#E0F7FA")
        scrollbar = Scrollbar(self.master.frame_contenido, orient=VERTICAL, command=canvas.yview)
        scrollable_frame = Frame(canvas, bg="#E0F7FA")

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")

        etiqueta_titulo = Label(scrollable_frame, text="Crear Historia Clínica", font=("Arial", 24, "bold"), bg="#00ACC1", fg="#FFFFFF")
        etiqueta_titulo.pack(pady=20)

        campos = [
            ("Cédula del Paciente:", 'cedula_entry', 1),
            ("Fecha (YYYY-MM-DD):", 'fecha_entry', 1),
            ("Motivo de Consulta:", 'motivo_consulta_entry', 5),
            ("Antecedentes:", 'antecedentes_entry', 5),
            ("Medicamentos Habituales:", 'medicamentos_entry', 5),
            ("Hábitos Nocivos:", 'habitos_entry', 5),
            ("Métodos Anticonceptivos:", 'anticonceptivos_entry', 5),
            ("Examen Físico:", 'examen_fisico_entry', 9),
            ("Diagnóstico Definitivo:", 'diagnostico_entry', 5)
        ]

        for texto, var, lines in campos:
            frame = Frame(scrollable_frame, bg="#E0F7FA")
            frame.pack(fill=X, pady=5)
            label = Label(frame, text=texto, bg="#E0F7FA", fg="#333333", font=("Arial", 14), width=25, anchor="w")
            label.pack(side=LEFT, padx=5)
            if texto.startswith("Fecha"):
                entry = DateEntry(frame, font=("Arial", 14), date_pattern='yyyy-mm-dd', bg="#FFFFFF", fg="#333333")
            elif lines == 1:
                entry = Entry(frame, font=("Arial", 14), bg="#FFFFFF", fg="#333333")
            else:
                entry = Text(frame, font=("Arial", 14), bg="#FFFFFF", fg="#333333", height=lines, wrap=WORD)
            setattr(self, var, entry)
            entry.pack(side=LEFT, fill=X, expand=YES, padx=5)

        frame_botones = Frame(scrollable_frame, bg="#E0F7FA")
        frame_botones.pack(pady=20)

        boton_guardar = Button(frame_botones, text="Guardar", command=self.guardar_historial, bg="#28A745", fg="white", font=("Arial", 16, "bold"), relief=RAISED, bd=5, width=15)
        boton_guardar.pack(side=LEFT, padx=10)

        boton_cancelar = Button(frame_botones, text="Cancelar", command=self.mostrar, bg="#DC3545", fg="white", font=("Arial", 16, "bold"), relief=RAISED, bd=5, width=15)
        boton_cancelar.pack(side=LEFT, padx=10)

        # Configurar peso para redimensionamiento
        self.master.frame_contenido.grid_rowconfigure(0, weight=1)
        self.master.frame_contenido.grid_columnconfigure(0, weight=1)

    def guardar_historial(self):
        cedula = self.cedula_entry.get()
        fecha = self.fecha_entry.get()
        motivo_consulta = self.motivo_consulta_entry.get("1.0", END).strip()
        antecedentes = self.antecedentes_entry.get("1.0", END).strip()
        medicamentos = self.medicamentos_entry.get("1.0", END).strip()
        habitos = self.habitos_entry.get("1.0", END).strip()
        anticonceptivos = self.anticonceptivos_entry.get("1.0", END).strip()
        examen_fisico = self.examen_fisico_entry.get("1.0", END).strip()
        diagnostico = self.diagnostico_entry.get("1.0", END).strip()

        paciente = obtener_paciente_por_cedula(cedula)
        if not paciente:
            messagebox.showerror("Error", "Paciente no encontrado")
            return
        paciente_id = paciente[0]

        try:
            agregar_historia_clinica(paciente_id, fecha, motivo_consulta, antecedentes, medicamentos, habitos, anticonceptivos, examen_fisico, diagnostico)
            messagebox.showinfo("Éxito", "Historia clínica creada correctamente")
            self.mostrar()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo crear la historia clínica: {e}")

    def eliminar_historial(self):
        selected_item = self.tree_historial_clinico.selection()
        if not selected_item:
            messagebox.showwarning("Advertencia", "Seleccione un historial para eliminar")
            return

        historia_id = self.tree_historial_clinico.item(selected_item)["values"][0]  # Suponiendo que la Cédula está en la primera columna
        respuesta = messagebox.askyesno("Confirmar Eliminación", "¿Está seguro que desea eliminar este historial?")
        if respuesta:
            try:
                eliminar_historial(historia_id)
                messagebox.showinfo("Éxito", "Historial eliminado correctamente")
                self.mostrar()  # Actualiza la lista de historiales para reflejar la eliminación
            except Exception as e:
                messagebox.showerror("Error", f"Error al eliminar el historial: {str(e)}")

    def exportar_historial(self):
        try:
            now = datetime.now()
            timestamp = now.strftime("%Y%m%d_%H%M%S")
            filename = f"exportaciones/historial_clinico_{timestamp}.xlsx"

            if not os.path.exists('exportaciones'):
                os.makedirs('exportaciones')

            exportar_historial_clinico_excel(filename)
            messagebox.showinfo("Éxito", f"Historial de historias clínicas exportado correctamente a {filename}")
        except Exception as e:
            messagebox.showerror("Error", f"Error al exportar las historias clínicas: {str(e)}")

    def exportar_historial_seleccionado(self):
        selected_item = self.tree_historial_clinico.selection()
        if not selected_item:
            messagebox.showwarning("Advertencia", "Seleccione un historial para exportar")
            return

        historia_id = self.tree_historial_clinico.item(selected_item)["values"][0]  # Suponiendo que la Cédula está en la primera columna
        try:
            historia = obtener_historial_clinico_por_cedula(historia_id)
            if historia:
                now = datetime.now()
                timestamp = now.strftime("%Y%m%d_%H%M%S")
                filename = f"exportaciones/historial_clinico_{historia_id}_{timestamp}.xlsx"

                if not os.path.exists('exportaciones'):
                    os.makedirs('exportaciones')

                df = pd.DataFrame(historia, columns=["ID", "Cédula", "Nombre", "Apellido", "Fecha", "Motivo Consulta", "Antecedentes", "Medicamentos", "Hábitos", "Métodos Anticonceptivos", "Examen Físico", "Diagnóstico"])
                df.to_excel(filename, index=False)
                messagebox.showinfo("Éxito", f"Historial clínico exportado correctamente a {filename}")
            else:
                messagebox.showinfo("Información", "No se encontró el historial clínico seleccionado")
        except Exception as e:
            messagebox.showerror("Error", f"Error al exportar el historial clínico: {str(e)}")

    def exportar_historial_seleccionado_word(self):
        selected_item = self.tree_historial_clinico.selection()
        if not selected_item:
            messagebox.showwarning("Advertencia", "Seleccione un historial para exportar")
            return

        historia_id = self.tree_historial_clinico.item(selected_item)["values"][0]  # Suponiendo que la Cédula está en la primera columna
        try:
            historia = obtener_historial_clinico_por_cedula(historia_id)
            if historia:
                now = datetime.now()
                timestamp = now.strftime("%Y%m%d_%H%M%S")
                filename = f"exportaciones/historial_clinico_{historia_id}_{timestamp}.docx"

                if not os.path.exists('exportaciones'):
                    os.makedirs('exportaciones')

                doc = Document()
                doc.add_heading('Historia Clínica', 0)

                for h in historia:
                    p = doc.add_paragraph()
                    p.add_run('Cédula: ').bold = True
                    p.add_run(h[1])

                    p = doc.add_paragraph()
                    p.add_run('Nombre: ').bold = True
                    p.add_run(f"{h[2]} {h[3]}")

                    p = doc.add_paragraph()
                    p.add_run('Fecha: ').bold = True
                    p.add_run(h[4])

                    p = doc.add_paragraph()
                    p.add_run('Motivo Consulta: ').bold = True
                    p.add_run(h[5])

                    p = doc.add_paragraph()
                    p.add_run('Antecedentes: ').bold = True
                    p.add_run(h[6])

                    p = doc.add_paragraph()
                    p.add_run('Medicamentos Habituales: ').bold = True
                    p.add_run(h[7])

                    p = doc.add_paragraph()
                    p.add_run('Hábitos Nocivos: ').bold = True
                    p.add_run(h[8])

                    p = doc.add_paragraph()
                    p.add_run('Métodos Anticonceptivos: ').bold = True
                    p.add_run(h[9])

                    p = doc.add_paragraph()
                    p.add_run('Examen Físico: ').bold = True
                    p.add_run(h[10])

                    p = doc.add_paragraph()
                    p.add_run('Diagnóstico Definitivo: ').bold = True
                    p.add_run(h[11])

                doc.save(filename)
                messagebox.showinfo("Éxito", f"Historial clínico exportado correctamente a {filename}")
            else:
                messagebox.showinfo("Información", "No se encontró el historial clínico seleccionado")
        except Exception as e:
            messagebox.showerror("Error", f"Error al exportar el historial clínico: {str(e)}")

if __name__ == "__main__":
    root = Tk()
    root.title("Sistema de Agendamiento de Citas")
    root.geometry("1200x800")
    app = PantallaHistoriasClinicas(root)
    app.mostrar()
    root.mainloop()
